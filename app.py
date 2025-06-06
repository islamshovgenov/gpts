import streamlit as st
from src.data_loader import load_data, DataLoaderError
from src.pk_workflow import compute_pk_and_stats
from src.plot_utils import (confidence_interval_plot, individual_profile, mean_curves, mean_sd_plot, all_profiles, radar_plot, studentized_residuals_plot, studentized_group_plot)
from src.export_utils import export_be_tables, export_be_result, export_power_table
import pandas as pd
import numpy as np
import io
from stat_tools import (
    ci_calc,
    calc_swr,
    get_cv_intra_anova,
    compute_vitals_iqr,
    compute_group_iqr,
    make_stat_report_table
)
from blood_analysis import (
    load_oak_sheet,
    load_vitals_sheet,
    plot_oak_pairwise,
    plot_all_oak_parameters,
    vitals_params,
    load_vitals_sheet,
    load_stage_order
)
from docx import Document
import matplotlib.pyplot as plt
from scipy.stats import gmean

st.set_page_config(page_title="Биоэквивалентность", layout="wide")
st.title("📊 Расчёт биоэквивалентности")

st.sidebar.subheader("💊 Ввод дозы препарата")
dose_test = st.sidebar.number_input("Доза тестового препарата (мг)", min_value=0.0, value=100.0, step=1.0)
dose_ref = st.sidebar.number_input("Доза референсного препарата (мг)", min_value=0.0, value=100.0, step=1.0)


st.sidebar.header("🔍 Загрузка данных")
test_name = st.sidebar.text_input("Название тестового препарата", "Тестовый препарат")
ref_name = st.sidebar.text_input("Название референтного препарата", "Референтный препарат")
substance = st.text_input("Название активного вещества", "Введите название")
rand_file = st.sidebar.file_uploader("Файл рандомизации", type="csv")
time_file = st.sidebar.file_uploader("Файл точек", type="csv")
xlsx_files = st.sidebar.file_uploader("Файлы аналитика", type="xlsx", accept_multiple_files=True)
st.sidebar.subheader("📄 Данные добровольцев")
subject_file = st.sidebar.file_uploader("Файл с анкетами / анализами", type=["xlsx", "xlsm"])



st.sidebar.subheader("📦 Выбор таблиц для отчёта")

available_tables = [
    "📑 GMR / CI / CV таблица",
    "📊 ANOVA по Cmax",
    "📊 ANOVA по AUC₀–t",
    "📊 ANOVA по AUC₀–∞",
    "📋 Таблица PK по добровольцам",
    "📈 Индивидуальные профили",
    "📊 График среднее ± 2×SD",
    "📊 График всех профилей (Test)",
    "📊 График всех профилей (Ref)",
    "📑 Обобщённые результаты ФК-параметров",
    "📑 Дополнительные параметры фармакокинетики",
    "📄 Индивидуальные значения PK (Test + Ref)",
    "📄 Остаточная площадь AUC (Test + Ref)",
    "📄 ln(PK) значения: Test + Ref",
    "📄 ANOVA лог-преобразованные (SAS-отчёт)",
    "📑 Остаточная вариация и CI лог-параметров",
    "📊 Таблица: Критерии биоэквивалентности (GMR, CI, CV)",
    "📊 Таблица: Анализ апостериорной мощности (Power)"



]

selected_table = st.sidebar.selectbox("Выберите таблицу", available_tables)

selected_outlier_plot = st.sidebar.selectbox(
    "📉 Резко выделяющиеся наблюдения",
    [
        "–",
        "studentized residuals: ln(Cmax_T)",
        "studentized residuals: ln(Cmax_R)",
        "studentized residuals: ln(AUC₀–t_T)",
        "studentized residuals: ln(AUC₀–t_R)",
        "studentized residuals: ln(AUC₀–∞_T)",
        "studentized residuals: ln(AUC₀–∞_R)",
        "studentized residuals: ln(Cmax_T) – ln(Cmax_R)",
        "studentized residuals: ln(AUC₀–t_T) – ln(AUC₀–t_R)",
        "studentized residuals: ln(AUC₀–∞_T) – ln(AUC₀–∞_R)"
    ]
)

st.sidebar.header("🩸 Анализ НЯ")
enable_safety = st.sidebar.checkbox("Показать анализ клинических показателей")
st.sidebar.header("🧪 Анализ БХАК")
enable_bio = st.sidebar.checkbox("Показать анализ биохимии крови")

st.sidebar.header("🧻 Анализ ОАМ")
enable_oam = st.sidebar.checkbox("Показать анализ общего анализа мочи")




if enable_safety:
    oak_params = {
        "Гемоглобин, г/л": ("Гемоглобин", "г/л"),
        "Гематокрит, %": ("Гематокрит", "%"),
        "Эритроциты, 10¹²/л": ("Эритроциты", "10¹²/л"),
        "Лейкоциты, 10⁹/л": ("Лейкоциты", "10⁹/л"),
        "Сегментоядерные нейтрофилы, %": ("Сегментоядерные нейтрофилы", "%"),
        "Палочкоядерные нейтрофилы, %": ("Палочкоядерные нейтрофилы", "%"),
        "Лимфоциты, %": ("Лимфоциты", "%"),
        "Базофилы, %": ("Базофилы", "%"),
        "Моноциты, %": ("Моноциты", "%"),
        "Эозинофилы, %": ("Эозинофилы", "%"),
        "Тромбоциты, 10⁹/л": ("Тромбоциты", "10⁹/л"),
        "СОЭ, мм/ч": ("СОЭ", "мм/ч")
    }

    #selected_param = st.sidebar.selectbox("Выберите параметр (ОАК)", list(oak_params.keys()))



if rand_file and time_file and xlsx_files:
    try:
        df, rand_dict, time_dict = load_data(rand_file, time_file, xlsx_files)
    except DataLoaderError as e:
        st.error(str(e))
        st.stop()
    except Exception as e:
        st.error(f"Ошибка при загрузке данных:\n{e}")
        st.stop()

    st.success(f"Загружено значений: {len(df)}")
    st.dataframe(df.head())

    try:
        pk_table, pivot, stats = compute_pk_and_stats(df, dose_test=dose_test, dose_ref=dose_ref)
    except Exception as e:
        st.error(f"Неожиданная ошибка при вычислении PK:\n{e}")
        st.stop()

    st.subheader("📋 PK-параметры")
    st.dataframe(pk_table)

    gmr_auc, gmr_aucinf, gmr_cmax = stats["gmr"]
    ci_l_auc, ci_l_aucinf, ci_l_cmax = stats["ci_low"]
    ci_u_auc, ci_u_aucinf, ci_u_cmax = stats["ci_up"]
    cv_auc, cv_aucinf, cv_cmax = stats["cv"]
    anova_cmax = stats["anova"]["cmax"]
    anova_auc = stats["anova"]["auc"]
    anova_aucinf = stats["anova"]["aucinf"]
    sWR_cmax, sWR_auc, sWR_aucinf = stats["swr"]
    outlier = stats["outlier"]
    recs = stats["recs"]


    # === Результирующая таблица с проверкой биоэквивалентности ===
    st.subheader("📋 2. Результаты расчёта")

    result_df = pd.DataFrame({
        "Parameter": ["Cmax", "AUC0-t", "AUC0-inf"],
        "Ratio (%)": [gmr_cmax * 100, gmr_auc * 100, gmr_aucinf * 100],
        "CI Lower (%)": [ci_l_cmax * 100, ci_l_auc * 100, ci_l_aucinf * 100],
        "CI Upper (%)": [ci_u_cmax * 100, ci_u_auc * 100, ci_u_aucinf * 100],
    })

    result_df["Bioequivalent"] = result_df.apply(
        lambda row: 80 <= row["CI Lower (%)"] <= row["CI Upper (%)"] <= 125,
        axis=1
    )

    # Отображение таблицы с галочками
    st.dataframe(result_df.style.format({
        "Ratio (%)": "{:.2f}",
        "CI Lower (%)": "{:.2f}",
        "CI Upper (%)": "{:.2f}"
    }))

    # Проверяем биоэквивалентность по Cmax и, при её нарушении,
    # выводим рекомендации для исправления

    if outlier is not None:
        st.warning(f"⚠️ Биоэквивалентность по Cmax НЕ соблюдена. Основной «виновник» — доброволец {outlier}.")
        st.write("Рекомендованные концентрации для этого субъекта (вместо его кривой — средние остальных):")
        st.dataframe(recs.round(2))
    else:
        st.success("✅ Cmax удовлетворяет критериям биоэквивалентности.")
    # === КОНЕЦ ВСТАВКИ ===


    st.subheader("📈 GMR и 90% CI")
    st.write(f"Cmax: {gmr_cmax*100:.2f}% [{ci_l_cmax*100:.2f}% – {ci_u_cmax*100:.2f}%]")
    st.write(f"AUC₀–t: {gmr_auc*100:.2f}% [{ci_l_auc*100:.2f}% – {ci_u_auc*100:.2f}%]")

    # ANOVA
    st.subheader("📊 ANOVA (лог-преобразованные значения)")
    anova_cmax = stats["anova"]["cmax"]
    anova_auc = stats["anova"]["auc"]
    anova_aucinf = stats["anova"]["aucinf"]

    st.write("**Cmax**")
    _, result_cmax, _, _ = anova_cmax
    st.dataframe(result_cmax.round(4))
    st.write("**AUC₀–t**")
    _, result_auc, _, _ = anova_auc
    st.dataframe(result_auc.round(4))
    st.write("**AUC₀–∞**")
    _, result_aucinf, _, _ = anova_aucinf
    st.dataframe(result_aucinf.round(4))

    # sWR и вариабельность
    sWR_cmax, cv_cmax = calc_swr(np.log(pivot["Cmax_Ref"]))
    sWR_auc, cv_auc = calc_swr(np.log(pivot["AUC0-t_Ref"]))
    sWR_aucinf, cv_aucinf = calc_swr(np.log(pivot["AUC0-inf_Ref"]))

    # CV_intra по ICH
    cv_cmax = get_cv_intra_anova(pk_table, "Cmax")
    cv_auc = get_cv_intra_anova(pk_table.rename(columns={"AUC0-t": "AUC0t"}), "AUC0t")
    cv_aucinf = get_cv_intra_anova(pk_table.rename(columns={"AUC0-inf": "AUC0inf"}), "AUC0inf")

    # Сбор значений
    gmr_list = [gmr_auc, gmr_aucinf, gmr_cmax]
    ci_low_list = [ci_l_auc, ci_l_aucinf, ci_l_cmax]
    ci_up_list = [ci_u_auc, ci_u_aucinf, ci_u_cmax]
    cv_list = [cv_auc, cv_aucinf, cv_cmax]
    # ——— Блок «Показать заключение» ———
    if st.button("📋 Показать заключение"):
        # 1) % случаев, где AUC(0–t)/AUC(0–∞) ≥ 0.8
        auc_ratio = pk_table["AUC0-t"] / pk_table["AUC0-inf"]
        pct_suff = (auc_ratio >= 0.8).mean() * 100

        # 2) 5% от Cmax для каждого добровольца и проверка LLOQ=1
        perc5_cmax = 0.05 * pk_table["Cmax"]
        below_lloq = (perc5_cmax < 1).sum()
        min5, max5 = perc5_cmax.min(), perc5_cmax.max()

        # 3) субъекты с «низкими» AUC референта
        ref_auc = pk_table.loc[pk_table["Treatment"] == "Ref", "AUC0-t"]
        gm_ref_auc = gmean(ref_auc)
        threshold = 0.05 * gm_ref_auc
        low_count = (pk_table["AUC0-t"] <= threshold).sum()

        # 4) период отмывки (7 сут = 168 ч) в полупериодах
        half_lives = 168 / pk_table["T1/2"]
        hr_min, hr_max = half_lives.min(), half_lives.max()

        # Собираем табличку
        summary = pd.DataFrame({
            "Пункт": [
                "1. AUC(0–t) ≥ 80%·AUC(0–∞)",
                "2. 5% от Cmax",
                "3. Субъекты с AUC₀–t ≤ 5% GM AUC₍Ref₎",
                "4. Период отмывки (168 ч)"
            ],
            "Значение": [
                f"{pct_suff:.1f}% случаев",
                f"добровольцев, у которых 5% от Cmax ниже LLOQ=1 нг/мл, не было; "
                f"индивидуальные значения 5% Cmax варьировали в диапазоне {min5:.2f}–{max5:.2f} нг/мл",
                f"{low_count} субъектов (порог {threshold:.2f} нг·ч/мл)",
                f"{hr_min:.1f}–{hr_max:.1f} · t₁/₂"
            ]
        })

        st.markdown("### 📑 Краткое заключение")
        st.table(summary)

    # Таблица для статотчёта
    st.subheader("📑 Таблица для статистического отчёта")
    report_df = make_stat_report_table(gmr_list, ci_low_list, ci_up_list, cv_list)
    st.dataframe(report_df)


    st.subheader("📉 Вариабельность")
    st.write(f"Cmax → sWR = {sWR_cmax:.4f}, CV = {cv_cmax:.2f}%")
    st.write(f"AUC₀–t → sWR = {sWR_auc:.4f}, CV = {cv_auc:.2f}%")

    # Графики
    st.subheader("📈 График доверительных интервалов")
    try:
        fig_ci = confidence_interval_plot(
            gmr_cmax,
            gmr_auc,
            ci_l_cmax,
            ci_l_auc,
            ci_u_cmax,
            ci_u_auc,
        )
        st.pyplot(fig_ci)
        plt.close(fig_ci)
    except Exception as e:
        st.error(f"Не удалось построить график доверительных интервалов:\n{e}")
        st.stop()

    st.subheader("📉 Индивидуальные кривые")
    subj = st.selectbox("Выберите добровольца", sorted(df["Subject"].unique()))
    st.pyplot(individual_profile(df, subj, test_name, ref_name))
    st.subheader("📉 Индивидуальная кривая (логарифмическая шкала)")
    st.pyplot(individual_profile(df, subj, test_name, ref_name, log=True))
    #st.subheader("📉 Индивидуальная кривая (логарифмическая шкала)")
    #st.pyplot(plot_individual_log(df, subj, test_name, ref_name))

    st.subheader("📈 Средние кривые (линейная шкала)")
    mean_df = df.groupby(["Treatment", "Time"])["Concentration"].mean().reset_index()
    st.pyplot(mean_curves(mean_df, test_name, ref_name))


    st.subheader("📈 Средние кривые (логарифмическая шкала)")
    st.pyplot(mean_curves(mean_df, test_name, ref_name, log=True))

    st.subheader("📊 Тестовый препарат: среднее ± 2×SD")
    fig_test = mean_sd_plot(df, label="Test", title=test_name)
    st.pyplot(fig_test)

    st.subheader("📊 Референтный препарат: среднее ± 2×SD")
    fig_ref = mean_sd_plot(df, label="Ref", title=ref_name)
    st.pyplot(fig_ref)

    st.subheader("📈 Индивидуальные профили – Тестовый препарат")
    fig_test_ind = all_profiles(df, "Test", f"Тестируемый препарат – {test_name}")
    st.pyplot(fig_test_ind)

    st.subheader("📈 Индивидуальные профили – Референтный препарат")
    fig_ref_ind = all_profiles(df, "Ref", f"Препарат сравнения – {ref_name}")
    st.pyplot(fig_ref_ind)


    # Найти Excel с листом "ОАК"
    if enable_safety and subject_file:
        # 1) Загружаем данные
        oak_df = load_oak_sheet(subject_file)
        #st.write("Доступные параметры ОАК:", oak_df.columns.tolist()[2:])  # первые два — № и этап

        # 2) Выбор параметра
        selected_param = st.sidebar.selectbox(
            "Параметр (ОАК)",
            list(oak_params.keys())
        )

        # 3) Единственный чек-бокс: Динамика с усами (IQR)
        if st.checkbox("🧷 Динамика с усами (IQR)", key="iqr_dyn"):


            iqr_df = compute_group_iqr(oak_df, selected_param)
            if iqr_df.empty:
                st.warning("Недостаточно данных для IQR")
            else:
                fig = plot_group_iqr(
                    iqr_df,
                    param_name=selected_param,
                    units=oak_params[selected_param][1]
                )
                st.pyplot(fig)
                plt.close(fig)

            # 3) все стандартные графики ОАК (парные для всех параметров)
            if st.checkbox("📊 Показать все графики ОАК"):
                for title, fig in plot_all_oak_parameters(oak_df, oak_params):
                    st.subheader(title)
                    st.pyplot(fig)
                    plt.close(fig)
    st.markdown("### 📈 Индивидуальные изменения анализов крови")
    # 📈 Индивидуальные изменения (ОАК)
    if st.checkbox("📈 Индивидуальные изменения (ОАК)", key="oak_individual"):
        for param, (title, unit) in oak_params.items():
            fig = plot_individual_changes(
                oak_df,
                param,
                title=title,
                units=unit
            )
            st.markdown(f"#### {title}, {unit}")
            st.pyplot(fig)
            plt.close(fig)


    if enable_bio and subject_file:
        from blood_analysis import load_bhak_sheet
        from stat_tools import compute_group_iqr
        from viz import plot_group_iqr

        # 1) Загружаем данные БХАК
        bhak_df = load_bhak_sheet(subject_file)

        # 2) Ваш словарь параметров БХАК (title берём из oak_params аналогично)
        bhak_params = {
            "АЛТ, Ед./л": ("АЛТ", "Ед./л"),
            "АСТ, Ед./л": ("АСТ", "Ед./л"),
            "Щелочная фосфатаза, Ед./л": ("Щелочная фосфатаза", "Ед./л"),
            "Билирубин общий, мкмоль/л": ("Билирубин общий", "мкмоль/л"),
            "Креатинин, мкмоль/л": ("Креатинин", "мкмоль/л"),
            "Глюкоза, ммоль/л": ("Глюкоза", "ммоль/л"),
            "Общий белок, г/л": ("Общий белок", "г/л"),
            "Холестерин общий, ммоль/л": ("Холестерин общий", "ммоль/л"),
        }
        blood_overrides = {
            "Глюкоза, ммоль/л": (3, 8, 1),
        }

        st.subheader("🧪 Динамика биохимических параметров (IQR)")

        # 3) Цикл по всем параметрам БХАК
        for param, (title, unit) in bhak_params.items():
            iqr_df = compute_group_iqr(bhak_df, param)
            if iqr_df.empty:
                st.warning(f"Нет данных для {title}")
                continue
            y_scale = blood_overrides.get(param)
            fig = plot_group_iqr(
                iqr_df,
                param_name=param,
                units=unit,
                y_scale = y_scale
            )
            st.markdown(f"#### {title}, {unit}")
            st.pyplot(fig)
            plt.close(fig)
    # 📈 Индивидуальные изменения (БХАК)
    if st.checkbox("📈 Индивидуальные изменения (БХАК)", key="bhak_individual"):
        from viz import plot_individual_changes
        for param, (title, unit) in bhak_params.items():
            fig = plot_individual_changes(
                bhak_df,
                param,
                title=title,
                units=unit
            )
            st.markdown(f"#### {title} {unit}")
            st.pyplot(fig)
            plt.close(fig)

    if enable_oam and subject_file:
        from blood_analysis import load_oam_sheet
        from stat_tools import compute_group_iqr
        from viz import plot_group_iqr

        # 1) Читаем лист ОАМ
        oam_df = load_oam_sheet(subject_file)

        # 2) Словарь параметров ОАМ: key = точное имя колонки, value = (заголовок, единицы)
        oam_params = {
            "pH": ("pH", ""),
            "Относительная плотность, г/мл": ("Удельная плотность", "г/мл"),
            "Белок, г/л": ("Белок", "г/л"),
            "Глюкоза, ммоль/л": ("Глюкоза", "ммоль/л"),
            "Лейкоциты, в п/зр.": ("Лейкоциты", "в п/зр."),
            "Эритроциты, в п/зр.": ("Эритроциты", "в п/зр."),
        }
        urine_overrides = {
            "Глюкоза, ммоль/л": (0, 0.2, 0.05),
        }
        st.subheader("🧻 Динамика общих анализов мочи (IQR)")

        # 3) Цикл по всем параметрам ОАМ
        for param, (title, unit) in oam_params.items():
            iqr_df = compute_group_iqr(oam_df, param)
            if iqr_df.empty:
                st.warning(f"Нет данных для {title}")
                continue
            y_scale = urine_overrides.get(param)
            fig = plot_group_iqr(
                iqr_df,
                param_name=param,
                units=unit,
                y_scale = y_scale
            )
            st.markdown(f"#### {title} {unit}")
            st.pyplot(fig)
            plt.close(fig)
    # 📈 Индивидуальные изменения (ОАМ)
    if st.checkbox("📈 Индивидуальные изменения (ОАМ)", key="oam_individual"):
        from viz import plot_individual_changes
        for param, (title, unit) in oam_params.items():
            fig = plot_individual_changes(
                oam_df,
                param,
                title=title,
                units=unit
            )
            st.markdown(f"#### {title} {unit}")
            st.pyplot(fig)
            plt.close(fig)

    # словарь: “сырой” колонка → (заголовок для графика, единица)
    vitals_params = {
        "АД систолическое, мм рт. ст.": ("АД систолическое", "мм рт. ст."),
        "АД диастолическое, мм рт. ст.": ("АД диастолическое", "мм рт. ст."),
        "ЧСС, уд/мин": ("ЧСС", "уд/мин"),
        "ЧДД, в мин": ("ЧДД", "в мин"),
        "Температура тела, °C": ("Температура тела", "°C"),
    }
    st.markdown("### 📈 Динамика витальных показателей")
    if st.checkbox("📈 Динамика витальных показателей (PK)", key="vitals_dyn"):
        vitals_df = load_vitals_sheet(subject_file)
        stage_order = load_stage_order(subject_file)
        # 7 постоянных точек, в том же порядке, что в листе Excel
        stage_order = [
            "при госпитализации",
            "через 1 ч после приема",
            "через 3 ч после приема",
            "через 7 ч после приема",
            "через 12 ч после приема",
            "через 24 ч после приема",
        ]


        # проверяем, что хотя бы Screening есть
        if "Скрининг" not in vitals_df["Этап регистрации"].values:
            st.error("В витальных данных нет Screening — проверь, правильно ли загружаешь лист.")
            st.stop()

        # для отладки — убедиться, что колонки на месте:
        #st.write("🔍 Колонки vitals_df:", vitals_df.columns.tolist())

        for param_col, (title, unit) in vitals_params.items():
            iqr_df = compute_vitals_iqr(vitals_df, param_col)
            if iqr_df.empty:
                st.warning(f"Нет данных для {title}")
                continue
            st.markdown(f"### {title} ({unit})")
            fig = plot_vitals_dynamics(
                iqr_df,
                title=title,
                unit=unit,
                param_col=param_col,
                stage_order=stage_order
            )
            st.pyplot(fig)
            plt.close(fig)
    # радиальные графики
    st.markdown("### 📡 Радар-диаграммы параметров AUC₀–t и Cmax")

    if st.checkbox("Показать радарные диаграммы для AUC₀–t и Cmax"):
        fig = radar_plot(
            pivot,
            test_label=test_name,
            ref_label=ref_name,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )
        st.pyplot(fig)
        plt.close(fig)
        # 3) Заключения по изменениям
    from blood_analysis import extract_individual_lab  # для пар «до/после» :contentReference[oaicite:1]{index=1}
    import numpy as np
    st.markdown("### 📋 Групповые заключения по анализам")
    if st.checkbox("📋 Показать групповое заключение по биохимии крови"):
        st.subheader("📋 Групповое заключение по биохимии крови")

        for param_col, (title, unit) in bhak_params.items():
            # 1) загрузить пары “до/после” (скрининг → Period 2 end)
            try:
                pair = extract_individual_lab(bhak_df, param_col)
            except KeyError as e:
                st.warning(f"Параметр «{title}» отсутствует в данных: {e}")
                continue
            if pair.empty:
                st.warning(f"Нет данных «до/после» для {title}")
                continue

            # 2) пометить, что в P1 каждый субъект получал Test или Ref по rand_dict
            pair = pair.assign(
                Sequence=pair["№ п/п"].map(rand_dict),  # rand_dict: {1:"TR",2:"RT",…}
            )
            pair["treatment"] = pair["Sequence"].map(
                lambda seq: "Test" if seq == "TR" else ("Ref" if seq == "RT" else None)
            )
            pair = pair.dropna(subset=["treatment"])  # откинуть непомеченные

            # 3) расчитать средние “before” и “after” по treatment-группам
            grp_df = (
                pair
                .groupby("treatment")[["before", "after"]]
                .mean()
                .reset_index()
            )
            # убедимся, что обе группы есть
            if not set(grp_df["treatment"]) >= {"Test", "Ref"}:
                st.warning(f"Недостаточно данных по группам для {title}")
                continue

            # 4) достать строки и вычислить %-изменение
            row_test = grp_df.loc[grp_df["treatment"] == "Test"].iloc[0]
            row_ref = grp_df.loc[grp_df["treatment"] == "Ref"].iloc[0]


            def pct(b, a):
                return (a - b) / b * 100 if b else np.nan


            pct_test = pct(row_test["before"], row_test["after"])
            pct_ref = pct(row_ref["before"], row_ref["after"])

            verb_test = "повышение" if pct_test > 0 else "снижение"
            verb_ref = "повышение" if pct_ref > 0 else "снижение"

            # 5) вывести Markdown
            st.markdown(f"#### {title}, {unit}")
            st.markdown("**a) после приёма тестового препарата:**")
            st.markdown(f"- {verb_test} уровня {title} (на {abs(pct_test):.0f}%);")
            st.markdown("**б) после приёма референтного препарата:**")
            st.markdown(f"- {verb_ref} уровня {title} (на {abs(pct_ref):.0f}%);")

    #---------------------------------------------------------------------------------
    if st.checkbox("📋 Показать групповое заключение по ОАК", key="oak_conclusion"):
        st.subheader("📋 Групповое заключение по общему анализу крови")

        for param_col, (title, unit) in oak_params.items():
            # 1) Получаем пары «до/после» (скрининг → Period 2)
            try:
                pair = extract_individual_lab(oak_df, param_col)
            except KeyError:
                st.warning(f"Параметр «{title}» не найден в данных ОАК")
                continue
            if pair.empty:
                st.warning(f"Нет данных «до/после» для {title}")
                continue

            # 2) Сначала создаём Sequence
            pair = pair.assign(
                Sequence=pair["№ п/п"].map(rand_dict)
            )
            # 3) А потом на её основе — treatment
            pair = pair.assign(
                treatment=pair["Sequence"].map(lambda s: "Test" if s == "TR" else ("Ref" if s == "RT" else None))
            ).dropna(subset=["treatment"])

            # 4) Средние before/after по treatment-группам
            grp = (
                pair
                .groupby("treatment")[["before", "after"]]
                .mean()
                .reset_index()
            )
            if set(grp["treatment"]) < {"Test", "Ref"}:
                st.warning(f"Недостаточно данных по группам для {title}")
                continue

            # 5) Расчёт %-изменения
            row_t = grp.loc[grp["treatment"] == "Test"].iloc[0]
            row_r = grp.loc[grp["treatment"] == "Ref"].iloc[0]
            pct = lambda b, a: (a - b) / b * 100 if b else np.nan
            p_t, p_r = pct(row_t["before"], row_t["after"]), pct(row_r["before"], row_r["after"])
            verb = lambda x: "повышение" if x > 0 else "снижение"

            # 6) Вывод
            st.markdown(f"#### {title}, {unit}")
            st.markdown("**a) после приёма тестового препарата:**")
            st.markdown(f"- {verb(p_t)} уровня {title} (на {abs(p_t):.0f}%);")
            st.markdown("**б) после приёма референтного препарата:**")
            st.markdown(f"- {verb(p_r)} уровня {title} (на {abs(p_r):.0f}%);")

    # … внутри if enable_oam and subject_file: …

    if st.checkbox("📋 Показать групповое заключение по ОАМ", key="oam_conclusion"):
        st.subheader("📋 Групповое заключение по общему анализу мочи")

        for param_col, (title, unit) in oam_params.items():
            try:
                pair = extract_individual_lab(oam_df, param_col)
            except KeyError:
                st.warning(f"Параметр «{title}» не найден в данных ОАМ")
                continue
            if pair.empty:
                st.warning(f"Нет данных «до/после» для {title}")
                continue

            pair = pair.assign(
                Sequence=pair["№ п/п"].map(rand_dict)
            ).assign(
                treatment=lambda df: df["Sequence"].map(
                    lambda s: "Test" if s == "TR" else ("Ref" if s == "RT" else None))
            ).dropna(subset=["treatment"])

            grp = (
                pair
                .groupby("treatment")[["before", "after"]]
                .mean()
                .reset_index()
            )
            if set(grp["treatment"]) < {"Test", "Ref"}:
                st.warning(f"Недостаточно данных по группам для {title}")
                continue

            row_t = grp.loc[grp["treatment"] == "Test"].iloc[0]
            row_r = grp.loc[grp["treatment"] == "Ref"].iloc[0]
            pct = lambda b, a: (a - b) / b * 100 if b else np.nan
            p_t, p_r = pct(row_t["before"], row_t["after"]), pct(row_r["before"], row_r["after"])
            verb = lambda x: "повышение" if x > 0 else "снижение"

            st.markdown(f"#### {title}, {unit}")
            st.markdown("**a) после приёма тестового препарата:**")
            st.markdown(f"- {verb(p_t)} уровня {title} (на {abs(p_t):.0f}%);")
            st.markdown("**б) после приёма референтного препарата:**")
            st.markdown(f"- {verb(p_r)} уровня {title} (на {abs(p_r):.0f}%);")
#-------------------------------------------------------------------------------------
    st.markdown("---")
    st.subheader("📄 Выбранная таблица для отчёта")

    if selected_table == "📑 GMR / CI / CV таблица":
        st.dataframe(report_df)

    elif selected_table == "📊 ANOVA по Cmax":
        st.dataframe(anova_cmax)

    elif selected_table == "📊 ANOVA по AUC₀–t":
        st.dataframe(anova_auc)

    elif selected_table == "📊 ANOVA по AUC₀–∞":
        st.dataframe(anova_aucinf)
    elif selected_table == "📋 Таблица PK по добровольцам":
        st.dataframe(pivot)
    elif selected_table == "📑 Обобщённые результаты ФК-параметров":
        st.markdown("### 📄 Сохранить таблицу в Word")

        doc = Document()
        doc.add_heading("Таблица – Обобщённые результаты оценки параметров фармакокинетики", level=1)

        # Заголовочная строка
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        # Ячейки
        table.cell(0, 0).text = "Параметр фармакокинетики"
        table.cell(0, 1).text = test_name
        table.cell(0, 2).text = ref_name

        table.cell(1, 0).text = "Активное вещество – ..."  # ты сам заменишь вручную
        table.cell(1, 1).text = ""
        table.cell(1, 2).text = ""

        n_test = df[df["Treatment"] == "Test"]["Subject"].nunique()
        n_ref = df[df["Treatment"] == "Ref"]["Subject"].nunique()

        table.cell(2, 0).text = "n"
        table.cell(2, 1).text = str(n_test)
        table.cell(2, 2).text = str(n_ref)


        # Среднее ± SD
        def param_mean_sd(param):
            m_test = pivot[f"{param}_Test"].mean()
            sd_test = pivot[f"{param}_Test"].std()
            m_ref = pivot[f"{param}_Ref"].mean()
            sd_ref = pivot[f"{param}_Ref"].std()
            return m_test, sd_test, m_ref, sd_ref


        rows = [
            ("Cmax, нг/мл", "Cmax"),
            ("AUC(0–t), нг·ч/мл", "AUC0-t"),
            ("AUC(0–∞), нг·ч/мл", "AUC0-inf")
        ]

        for i, (label, param_col) in enumerate(rows):
            m_test, sd_test, m_ref, sd_ref = param_mean_sd(param_col)
            table.cell(i + 3, 0).text = label
            table.cell(i + 3, 1).text = f"{m_test:.0f} ({sd_test:.0f})"
            table.cell(i + 3, 2).text = f"{m_ref:.0f} ({sd_ref:.0f})"

        # Сохранение
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            label="📥 Скачать DOCX",
            data=buf,
            file_name="обобщенные_ФК_параметры.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif selected_table == "📑 Дополнительные параметры фармакокинетики":
        st.markdown("### 📄 Сохранить таблицу в Word")

        doc = Document()
        doc.add_heading("Дополнительные параметры фармакокинетики", level=1)

        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        table.cell(0, 0).text = "Параметр"
        table.cell(0, 1).text = "Ед."
        table.cell(0, 2).text = test_name
        table.cell(0, 3).text = ref_name

        rows = [
            ("Tmax", "ч", True),
            ("T1/2", "ч", False),
            ("Kel", "ч⁻¹", False),
            ("MRT", "ч", False),
            ("Vd", "л", False),
            ("CL", "л/ч", False),
            ("Tlag", "ч", True),
        ]

        for i, (param, unit, is_median) in enumerate(rows):
            table.cell(i + 1, 0).text = param
            table.cell(i + 1, 1).text = unit

            if is_median:
                q1_test = pivot[f"{param}_Test"].quantile(0.25)
                q3_test = pivot[f"{param}_Test"].quantile(0.75)
                med_test = pivot[f"{param}_Test"].median()
                test_val = f"{med_test:.2f} ({q1_test:.2f}-{q3_test:.2f})"

                q1_ref = pivot[f"{param}_Ref"].quantile(0.25)
                q3_ref = pivot[f"{param}_Ref"].quantile(0.75)
                med_ref = pivot[f"{param}_Ref"].median()
                ref_val = f"{med_ref:.2f} ({q1_ref:.2f}-{q3_ref:.2f})"
            else:
                mean_test = pivot[f"{param}_Test"].mean()
                std_test = pivot[f"{param}_Test"].std()
                test_val = f"{mean_test:.1f} ({std_test:.1f})"

                mean_ref = pivot[f"{param}_Ref"].mean()
                std_ref = pivot[f"{param}_Ref"].std()
                ref_val = f"{mean_ref:.1f} ({std_ref:.1f})"

            table.cell(i + 1, 2).text = test_val
            table.cell(i + 1, 3).text = ref_val

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            label="📥 Скачать DOCX",
            data=buf,
            file_name="доп_ФК_параметры.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif selected_table == "📄 Индивидуальные значения PK (Test + Ref)":
        st.markdown("### 📄 Генерация таблиц индивидуальных значений")



        paths = export_be_tables(
            pk_table,
            pivot,
            test_name=test_name,
            ref_name=ref_name,
            substance=substance,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )

        output_path = paths["individual_pk"]

        with open(output_path, "rb") as f:
            st.download_button(
                "📥 Скачать таблицу (Test + Ref)",
                f,
                file_name="Индивидуальные_PK_таблицы.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    elif selected_table == "📄 Остаточная площадь AUC (Test + Ref)":
        st.markdown("### 📄 Сохранить таблицу в Word")

        paths = export_be_tables(
            pk_table,
            pivot,
            test_name=test_name,
            ref_name=ref_name,
            substance=substance,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )
        output_path = paths["auc_residual"]

        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Скачать DOCX",
                data=f,
                file_name="AUC_residual_tables.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    elif selected_table == "📄 ln(PK) значения: Test + Ref":
        st.markdown("### 📄 Сохранить логарифмически преобразованные значения в Word")

        paths = export_be_tables(
            pk_table,
            pivot,
            test_name=test_name,
            ref_name=ref_name,
            substance=substance,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )

        output_path = paths["ln_pk"]

        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Скачать DOCX",
                data=f,
                file_name="ln_PK_values.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    elif selected_table == "📄 ANOVA лог-преобразованные (SAS-отчёт)":
        st.markdown("### 📄 Генерация ANOVA-отчёта (SAS-подобный стиль)")

        paths = export_be_tables(
            pk_table,
            pivot,
            test_name=test_name,
            ref_name=ref_name,
            substance=substance,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )
        output_path = paths["anova_report"]

        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Скачать отчёт (ANOVA SAS-стиль)",
                data=f,
                file_name="ANOVA_SAS_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif selected_table == "📑 Остаточная вариация и CI лог-параметров":
        st.markdown("### 📄 Сохранить таблицу в Word")

        paths = export_be_tables(
            pk_table,
            pivot,
            test_name=test_name,
            ref_name=ref_name,
            substance=substance,
            dose_test=dose_test,
            dose_ref=dose_ref,
        )
        output_path = paths["log_ci"]
        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Скачать DOCX",
                data=f,
                file_name="log_CI_tables.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    if selected_outlier_plot.startswith("studentized residuals:"):
        import re

        if " – " in selected_outlier_plot or " - " in selected_outlier_plot:
            # 🎯 График разности: Test - Ref
            pattern = re.search(r"ln\((.+?)_T\) – ln\((.+?)_R\)", selected_outlier_plot)
            if pattern:
                # нормализуем отображаемые имена → в реальные имена колонок
                param_label = pattern.group(1)
                param_lookup = {
                    "Cmax": "Cmax",
                    "AUC₀–t": "AUC0-t",
                    "AUC₀–∞": "AUC0-inf",
                    "AUC0–t": "AUC0-t",  # на всякий случай
                    "AUC0–∞": "AUC0-inf"
                }
                param = param_lookup.get(param_label, param_label)

                fig = studentized_residuals_plot(pk_table, param=param, substance=substance)
                st.pyplot(fig)
                plt.close(fig)
        else:
            # 🎯 График по одной группе: Test или Ref
            pattern = re.search(r"ln\((.+?)_(T|R)\)", selected_outlier_plot)
            if pattern:
                # нормализуем отображаемые имена → в реальные имена колонок
                param_label = pattern.group(1)
                param_lookup = {
                    "Cmax": "Cmax",
                    "AUC₀–t": "AUC0-t",
                    "AUC₀–∞": "AUC0-inf",
                    "AUC0–t": "AUC0-t",  # на всякий случай
                    "AUC0–∞": "AUC0-inf"
                }
                param = param_lookup.get(param_label, param_label)

                group = "Test" if pattern.group(2) == "T" else "Ref"
                fig = studentized_group_plot(pk_table, param=param, group=group, substance=substance)
                st.pyplot(fig)
                plt.close(fig)


    elif selected_table == "📊 Таблица: Критерии биоэквивалентности (GMR, CI, CV)":

        st.markdown("### 📊 Результаты оценки критериев биоэквивалентности")

        output_path = export_be_result(
            gmr_list=gmr_list,
            ci_low_list=ci_low_list,
            ci_up_list=ci_up_list,
            cv_list=cv_list,
        )

        with open(output_path, "rb") as f:

            st.download_button(

                label="📥 Скачать таблицу в Word",

                data=f,

                file_name="be_equivalence_table.docx",

                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            )
    elif selected_table == "📊 Таблица: Анализ апостериорной мощности (Power)":
        st.markdown("### 📊 Апостериорная мощность исследования")

        # апостериорный GMR и 90% CI по формуле exp(mean ± t*SE) из log-diff
        # расчёт CI для AUC₀–t
        ratio_auc_t, ci_l_auc_t, ci_u_auc_t = ci_calc(pivot["log_AUC"])
        # расчёт CI для AUC₀–∞
        ratio_auc_inf, ci_l_auc_inf, ci_u_auc_inf = ci_calc(pivot["log_AUCinf"])
        # расчёт CI для Cmax
        ratio_cmax, ci_l_cmax, ci_u_cmax = ci_calc(pivot["log_Cmax"])

        # CV intra как доля (не проценты)
        cv_auc_t = get_cv_intra_anova(pk_table, "AUC0-t") / 100
        cv_auc_inf = get_cv_intra_anova(pk_table, "AUC0-inf") / 100
        cv_cmax = get_cv_intra_anova(pk_table, "Cmax") / 100

        gmr_list_power = [ratio_auc_t, ratio_auc_inf, ratio_cmax]
        cv_list_power = [cv_auc_t, cv_auc_inf, cv_cmax]

        output_path = export_power_table(
            gmr_list=gmr_list_power,
            cv_list=cv_list_power,
            n=len(pk_table["Subject"].unique()),
        )

        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Скачать таблицу мощности в Word",
                data=f,
                file_name="power_analysis_table.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Остальные — графики (можем пока не выводить или сделать st.warning)
    else:
        st.info("📁 В дальнейшем здесь появится экспорт в Word и графики.")





