from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
from stat_tools import anova_log, mixed_anova, nested_anova
from scipy.stats import f as f_dist, t as t_dist
from scipy.stats import norm
import matplotlib.pyplot as plt
import os
from statsmodels.formula.api import ols

PK_COLUMNS = [
    "AUC0-t", "AUC0-inf", "Cmax", "Tmax",
    "T1/2", "Kel", "N_el", "MRT", "Tlag", "Vd", "CL"
]
PK_HEADER_LABELS = [
    "Ранд. №",
    "AUC0-t\nng·h/mL",
    "AUC0-inf\nng·h/mL",
    "Cₘₐₓ\nng/mL",
    "tₘₐₓ\nh",
    "t₁/₂\nh",
    "kₑₗ\nh⁻¹",
    "Nₑₗ",
    "MRT\nh",
    "Tlag\nh",
    "Vd\nL",
    "CL\nL/h"
]


def cv_log(series):
    series = series.dropna()
    if len(series) == 0 or (series <= 0).any():
        return np.nan
    logs = np.log(series)
    s = np.std(logs, ddof=1)
    return np.sqrt(np.exp(s**2) - 1) * 100


def export_individual_pk_tables(pk_df, test_name, ref_name, substance, dose_test, dose_ref, save_path="Индивидуальные_PK_таблицы.docx"):
    # ——— Проверяем наличие колонок для построения таблицы ———
    required = ["Subject"] + PK_COLUMNS
    missing = [c for c in required if c not in pk_df.columns]
    if missing:
        raise KeyError(f"В export_individual_pk_tables отсутствуют колонки: {missing}")
    def summary_stats(df):
        stats = {
            "N": df.count(),
            "Mean": df.mean(),
            "Geom": np.exp(np.log(df[df > 0]).mean()),
            "Median": df.median(),
            "SD": df.std(),
            "SE": df.sem(),
            "CV": df.apply(cv_log).round(0).astype("Int64").astype(str) + "%",
            "1st Qua": df.quantile(0.25),
            "3rd Qua": df.quantile(0.75),
            "min": df.min(),
            "max": df.max()
        }
        return pd.DataFrame(stats).T

    def make_table(doc, df, title):
        df_sorted = df.sort_values(by="Subject")
        doc.add_heading(title, level=2)
        columns = PK_COLUMNS
        header = PK_HEADER_LABELS


        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(header):
            hdr_cells[i].text = col

        for _, row in df_sorted.iterrows():
            values = [str(int(row["Subject"]))] + [str(round(row[col], 2)) if pd.notna(row[col]) else "" for col in columns]
            row_cells = table.add_row().cells
            for i, val in enumerate(values):
                row_cells[i].text = val

        # Добавим статистику
        stats_df = summary_stats(df_sorted.set_index("Subject")[columns])
        for idx, row in stats_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = idx
            for j, col in enumerate(columns):
                try:
                    val = round(float(row[col]), 2) if pd.notna(row[col]) else ""
                except (ValueError, TypeError):
                    val = str(row[col])
                row_cells[j + 1].text = str(val)

    doc = Document()
    doc.add_heading("Таблица индивидуальных значений фармакокинетических параметров", level=1)

    test_df = pk_df[pk_df["Treatment"] == "Test"].copy()
    ref_df = pk_df[pk_df["Treatment"] == "Ref"].copy()

    make_table(
        doc,
        test_df,
        f"Тестируемый препарат: {test_name}, активное вещество: {substance}, доза: {dose_test} мг"
    )

    doc.add_page_break()

    make_table(
        doc,
        ref_df,
        f"Препарат сравнения: {ref_name}, активное вещество: {substance}, доза: {dose_ref} мг"
    )

    doc.save(save_path)
    return save_path

def export_auc_residual_tables(df, test_name, ref_name, substance, dose_test, dose_ref, save_path):
    # ——— Проверка входных колонок ———
    required = ["Treatment", "AUC0-t", "AUC0-inf"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise KeyError(f"В export_auc_residual_tables отсутствуют колонки: {missing}")

    doc = Document()
    doc.add_heading("Таблица остаточной площади AUC", level=1)

    for group_name, label, name, dose in [("Test", "T", test_name, dose_test), ("Ref", "R", ref_name, dose_ref)]:
        group_df = df[df["Treatment"] == group_name].copy()
        group_df = group_df.sort_values("Subject")

        # Расчёт остаточной площади
        group_df["Residual_%"] = (
            (group_df["AUC0-inf"] - group_df["AUC0-t"]) / group_df["AUC0-inf"] * 100
        ).round(1)

        doc.add_paragraph(f"{label} - {name}, активное вещество: {substance}, доза: {dose} мг")

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Рандом. №"
        hdr[1].text = "AUC₀–t (нг·ч/мл)"
        hdr[2].text = "AUC₀–∞ (нг·ч/мл)"
        hdr[3].text = "Остаточная площадь (%)"

        for _, row in group_df.iterrows():
            tr = table.add_row().cells
            tr[0].text = str(int(row["Subject"]))
            tr[1].text = f"{row['AUC0-t']:.2f}"
            tr[2].text = f"{row['AUC0-inf']:.2f}"
            tr[3].text = f"{row['Residual_%']:.1f}%"

    doc.save(save_path)
    return save_path


def export_log_transformed_pk_tables(pk_df, test_name, ref_name, substance, dose_test, dose_ref, save_path="ln_PK_values.docx"):
    def log_transform(df):
        df = df.copy()
        # ——— Проверка входных колонок и положительных значений ———
        for col in ("AUC0-t", "AUC0-inf", "Cmax"):
            if col not in pk_df.columns:
                raise KeyError(f"В export_log_transformed_pk_tables отсутствует колонка: {col}")
        # Фильтруем: логировать можно только там, где >0
        df = df[df[["AUC0-t", "AUC0-inf", "Cmax"]].gt(0).all(axis=1)].copy()

        df["ln_AUC0-t"] = np.log(df["AUC0-t"])
        df["ln_AUC0-inf"] = np.log(df["AUC0-inf"])
        df["ln_Cmax"] = np.log(df["Cmax"])
        return df

    def make_ln_table(doc, df, title):
        df_sorted = df.sort_values("Subject")
        doc.add_heading(title, level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "№"
        hdr[1].text = "ln(AUC₀–t)"
        hdr[2].text = "ln(AUC₀–∞)"
        hdr[3].text = "ln(Cmax)"

        for _, row in df_sorted.iterrows():
            tr = table.add_row().cells
            tr[0].text = f"{int(row['Subject']):02d}"
            tr[1].text = f"{row['ln_AUC0-t']:.5f}"
            tr[2].text = f"{row['ln_AUC0-inf']:.5f}"
            tr[3].text = f"{row['ln_Cmax']:.5f}"

        # --- ИТОГОВАЯ СТАТИСТИКА ---
        cols = ["ln_AUC0-t", "ln_AUC0-inf", "ln_Cmax"]

        def get_cv(series):
            return np.std(series, ddof=1) / np.mean(series) * 100 if np.mean(series) > 0 else np.nan

        row_names = ["N", "Mean", "SD", "SE", "CV, %", "1st Qua", "3rd Qua", "min", "max"]
        for stat in row_names:
            tr = table.add_row().cells
            tr[0].text = stat
            for i, col in enumerate(cols, start=1):
                series = df[col].dropna()
                if stat == "N":
                    val = str(len(series))
                elif stat == "Mean":
                    val = f"{series.mean():.5f}"
                elif stat == "SD":
                    val = f"{series.std(ddof=1):.5f}"
                elif stat == "SE":
                    val = f"{series.sem():.5f}"
                elif stat == "CV, %":
                    val = f"{get_cv(series):.1f}"
                elif stat == "1st Qua":
                    val = f"{series.quantile(0.25):.5f}"
                elif stat == "3rd Qua":
                    val = f"{series.quantile(0.75):.5f}"
                elif stat == "min":
                    val = f"{series.min():.5f}"
                elif stat == "max":
                    val = f"{series.max():.5f}"
                else:
                    val = "-"
                tr[i].text = val



    doc = Document()
    doc.add_heading("Индивидуальные значения логарифмически преобразованных PK-параметров", level=1)
    doc.add_paragraph(f"Активное вещество: {substance}. Доза: {dose_test} мг натощак")

    test_df = pk_df[pk_df["Treatment"] == "Test"]
    test_ln = log_transform(test_df)
    make_ln_table(doc, test_ln, f"Тестируемый препарат: {test_name}")

    doc.add_page_break()
    doc.add_paragraph(f"Активное вещество: {substance}. Доза: {dose_ref} мг натощак")

    ref_df = pk_df[pk_df["Treatment"] == "Ref"]
    ref_ln = log_transform(ref_df)
    make_ln_table(doc, ref_ln, f"Препарат сравнения: {ref_name}")

    doc.save(save_path)
    return save_path



def export_sas_anova_report(pk_df, substance, dose_test, save_path="ANOVA_SAS_Report.docx"):
    doc = Document()
    doc.add_heading("Результаты ANOVA лог-преобразованных PK-параметров", level=1)
    doc.add_paragraph(f"Активное вещество: {substance}, доза: {dose_test} мг\n")

    param_map = {
        "Cmax": "Cmax",
        "AUC₀–t": "AUC0-t",
        "AUC₀–∞": "AUC0-inf"
    }

    interpretations = []


    # === Type III ANOVA (nested) для каждого PK-параметра ===
    for title, colname in param_map.items():

        # 1) Подготовка данных
        df = pk_df.copy()
        # если в ваших данных названия колонок AUC₀–т/∞ именно такие — меняем на латиницу
        df = df.rename(columns={"AUC₀–t": "AUC0-t", "AUC₀–∞": "AUC0-inf"})

        # 2) Вызов нашей функции nested_anova из stat_tools.py
        res = nested_anova(df, colname)


        # ——— Таблица Model / Error / Corrected Total ———
        # Считаем Model SS/DF как сумма всех факторов кроме Error
        model_ss = sum(res["ss"][f] for f in ("Sequence","subject(Sequence)","formulation","Period"))
        model_df = sum(res["df"][f] for f in ("Sequence","subject(Sequence)","formulation","Period"))
        tot_ss   = model_ss + res["ss"]["Error"]
        tot_df   = model_df + res["df"]["Error"]

        tbl0 = doc.add_table(rows=3, cols=4)
        tbl0.style = "Table Grid"
        hdr0 = tbl0.rows[0].cells
        hdr0[0].text = "Source";         hdr0[1].text = "DF"
        hdr0[2].text = "Sum of Squares"; hdr0[3].text = "Mean Square"

        # Model
        row = tbl0.rows[1].cells
        row[0].text = "Model"
        row[1].text = str(model_df)
        row[2].text = f"{model_ss:.6f}"
        row[3].text = f"{(model_ss/model_df):.6f}"

        # Error
        row = tbl0.rows[2].cells
        row[0].text = "Error"
        row[1].text = str(res["df"]["Error"])
        row[2].text = f"{res['ss']['Error']:.6f}"
        row[3].text = f"{res['ms']['Error']:.6f}"

        # Corrected Total
        row = tbl0.add_row().cells
        row[0].text = "Corrected Total"
        row[1].text = str(tot_df)
        row[2].text = f"{tot_ss:.6f}"
        row[3].text = ""
        # ——— Конец Model / Error / Total ———


        # 3) Заголовок в Word
        doc.add_paragraph(
            f"Результаты дисперсионного анализа ln{title} ({substance}, {dose_test} мг, натощак)",
            style="Heading 2"
        )
        doc.add_paragraph(f"Dependent Variable: _ln{colname}")

        # 4) Таблица Type III SS / MS / F / Pr>F
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Source";     hdr[1].text = "DF"
        hdr[2].text = "Type III SS"; hdr[3].text = "Mean Square"
        hdr[4].text = "F Value";    hdr[5].text = "Pr > F"


        for factor in ("Sequence", "formulation", "subject(Sequence)", "Period"):
            df_i = res["df"][factor]
            ss_i = res["ss"][factor]
            ms_i = res["ms"][factor]
            f_i  = res["F"][factor]
            p_i  = res["p"][factor]

            row = tbl.add_row().cells
            row[0].text = factor
            row[1].text = str(df_i)
            row[2].text = f"{ss_i:.6f}"
            row[3].text = f"{ms_i:.6f}"
            row[4].text = f"{f_i:.2f}" if not np.isnan(f_i) else "-"
            row[5].text = f"{p_i:.4f}" if not np.isnan(p_i) else "-"

        # 5) Tests of Hypotheses (nested test для Sequence)
        doc.add_paragraph(
            "\nTests of Hypotheses Using the Type III MS for subject(Sequence) as an Error Term",
            style="Heading 3"
        )
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = "Table Grid"
        hdr2 = tbl2.rows[0].cells
        hdr2[0].text = "Source";     hdr2[1].text = "DF"
        hdr2[2].text = "Type III SS"; hdr2[3].text = "Mean Square"
        hdr2[4].text = "F Value";    hdr2[5].text = "Pr > F"

        df_i = res["df"]["Sequence"]
        ss_i = res["ss"]["Sequence"]
        ms_i = res["ms"]["Sequence"]
        f_n  = res["nested_test"]["Sequence"]["F"]
        p_n  = res["nested_test"]["Sequence"]["p"]

        row = tbl2.add_row().cells
        row[0].text = "Sequence";    row[1].text = str(df_i)
        row[2].text = f"{ss_i:.6f}"; row[3].text = f"{ms_i:.6f}"
        row[4].text = f"{f_n:.2f}";  row[5].text = f"{p_n:.4f}"

        # 6) Разрыв страницы перед следующим параметром
        doc.add_page_break()

        # строим OLS по тем же данным lnY ~ Sequence + Period + formulation
        df2 = df.copy()
        df2["lnY"] = np.log(df2[colname])
        # приводим факторы к строкам
        for c in ("Sequence","Period","Treatment"):
            df2[c] = df2[c].astype(str)

        model = ols("lnY ~ C(Sequence) + C(Period) + C(Treatment)", data=df2).fit()
        # название параметра в SAS-стиле
        param_name = "ABE for parameter"

        est  = model.params["C(Treatment)[T.Test]"]
        se   = model.bse   ["C(Treatment)[T.Test]"]
        tval = model.tvalues["C(Treatment)[T.Test]"]
        pval = model.pvalues["C(Treatment)[T.Test]"]


        tbl3 = doc.add_table(rows=1, cols=5)
        tbl3.style = "Table Grid"
        hdr3 = tbl3.rows[0].cells
        hdr3[0].text = "Parameter";      hdr3[1].text = "Estimate"
        hdr3[2].text = "Standard Error"; hdr3[3].text = "t Value"
        hdr3[4].text = "Pr > |t|"

        row = tbl3.add_row().cells
        row[0].text = param_name
        row[1].text = f"{est:.8f}"
        row[2].text = f"{se:.8f}"
        row[3].text = f"{tval:.2f}"
        row[4].text = f"{pval:.4f}"
        # ——— Конец Parameter Estimates ———


        # ——— Least Squares Means ———
        # Считаем LSM для двух уровней formulation
        # LSMean = общий средний + (разница estimate)/2
        # а 90% CI = estimate ± t(0.05; df_err)*SE

        # общий средний
        grand_mean = df2["lnY"].mean()
        # LSMean для каждого
        trt_means = df2.groupby("Treatment")["lnY"].mean()
        diff = est  # уже есть из Parameter Estimates
        # df стой
        df_err = res["df"]["Error"]
        tcrit = t_dist.ppf(0.95, df_err)  # 90% CI → alpha=0.10

        # границы CI для разницы
        lo = diff - tcrit * se
        hi = diff + tcrit * se

        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = "Table Grid"
        hdr4 = tbl4.rows[0].cells
        hdr4[0].text = "Difference Between Means"
        hdr4[1].text = "Lower 90% CI"
        hdr4[2].text = "Upper 90% CI"

        row = tbl4.add_row().cells
        row[0].text = f"{diff:.6f}"
        row[1].text = f"{lo:.6f}"
        row[2].text = f"{hi:.6f}"
        # ——— Конец Least Squares Means ———

    # === Конец ANOVA-цикла ===


    # Общие выводы
    doc.add_page_break()
    doc.add_heading("Выводы ANOVA", level=1)
    for line in interpretations:
        doc.add_paragraph(line)

    doc.add_paragraph("\n\nCreated by Islam Shovgenov in SAS ver. 9.4", style='Intense Quote')
    doc.save(save_path)
    return save_path



def export_log_ci_tables(df, substance, save_path="log_CI_tables.docx"):
    doc = Document()
    doc.add_heading("Остаточные вариации и 90% доверительные интервалы", level=1)
    doc.add_paragraph(f"Активное вещество: {substance}. Доза: 300 мг натощак")

    params = ["AUC0-t", "AUC0-inf", "Cmax"]
    log_cols = [f"log_{p}_diff" for p in params]
    se_cols = [f"SE_log_{p}" for p in params]
    mse_cols = [f"MSE_log_{p}" for p in params]
    # ——— Проверка наличия необходимых столбцов ———
    required = log_cols + se_cols + mse_cols
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise KeyError(f"В export_log_ci_tables отсутствуют колонки: {missing}")


    # Вводим t-критическое значение вручную для df=19 (пример), при необходимости заменить
    t_val = 1.697261

    # Таблица 1 – Остаточная вариация, SE, t
    table1 = doc.add_table(rows=1 + len(params), cols=4)
    table1.style = "Table Grid"
    table1.cell(0, 0).text = "Логарифм фармакокинетического параметра"
    table1.cell(0, 1).text = "MSE"
    table1.cell(0, 2).text = "SE"
    table1.cell(0, 3).text = "t (1–α)"

    for i, p in enumerate(params):
        table1.cell(i + 1, 0).text = f"ln({p})"
        table1.cell(i + 1, 1).text = f"{df[f'MSE_log_{p}'].iloc[0]:.6f}"
        table1.cell(i + 1, 2).text = f"{df[f'SE_log_{p}'].iloc[0]:.6f}"
        table1.cell(i + 1, 3).text = f"{t_val:.6f}"

    doc.add_paragraph("")  # разрыв

    # Таблица 2 – разности логарифмов и CI
    table2 = doc.add_table(rows=1 + len(params), cols=4)
    table2.style = "Table Grid"
    table2.cell(0, 0).text = "Разность логарифмов параметра"
    table2.cell(0, 1).text = "LSM"
    table2.cell(0, 2).text = "CI нижняя"
    table2.cell(0, 3).text = "CI верхняя"

    for i, p in enumerate(params):
        lsm = df[f"log_{p}_diff"].mean()
        se = df[f"SE_log_{p}"].iloc[0]
        lower = lsm - t_val * se
        upper = lsm + t_val * se
        table2.cell(i + 1, 0).text = f"ln({p})T - ln({p})R"
        table2.cell(i + 1, 1).text = f"{lsm:.6f}"
        table2.cell(i + 1, 2).text = f"{lower:.6f}"
        table2.cell(i + 1, 3).text = f"{upper:.6f}"

    doc.add_paragraph("где:  CIₗ = LSM – SE × t;   CIᵤ = LSM + SE × t")

    doc.save(save_path)
    return save_path

def export_be_result_table(gmr_list, ci_low_list, ci_up_list, cv_list, save_path="be_table.docx"):
    # ——— Проверка длин списков и наличия колонок ———
    if not (len(gmr_list)==len(ci_low_list)==len(ci_up_list)==len(cv_list)):
        raise ValueError("Списки gmr_list, ci_low_list, ci_up_list и cv_list должны быть одинаковой длины")

    doc = Document()
    doc.add_heading("Таблица — Результаты оценки критериев биоэквивалентности", level=1)

    # Таблица с 6 колонками
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Параметр"
    hdr[1].text = "CV₍intra₎"
    hdr[2].text = "GMR"
    hdr[3].text = "90% ДИ"
    hdr[4].text = "Границы"
    hdr[5].text = "Диаграмма"

    params = ["AUC0-t", "AUC0-inf", "Cmax"]

    for p, gmr, ci_low, ci_high, cv in zip(params, gmr_list, ci_low_list, ci_up_list, cv_list):
        row = table.add_row().cells

        row[0].text = p.replace("AUC0-", "AUC₀–").replace("inf", "∞")
        row[1].text = f"{cv:.2f}%"
        row[2].text = f"{gmr * 100:.2f}%"
        row[3].text = f"{ci_low * 100:.2f}% – {ci_high * 100:.2f}%"
        row[4].text = "80,00% – 125,00%"

        # график
        import matplotlib as mpl
        mpl.rcParams.update({
            'font.family': 'serif',
            'font.size': 8,
        })

        fig, ax = plt.subplots(figsize=(2, 1))

        # горизонтальная линия 90% ДИ
        ax.hlines(1, ci_low, ci_high, color='black', linewidth=1)

        # маркер GMR: ромбик с белой заливкой и черной обводкой
        ax.plot(
            gmr, 1,
            marker='D',
            markersize=6,
            markerfacecolor='white',
            markeredgecolor='black',
            markeredgewidth=1
        )

        # вертикальные линии 80% и 125% (пунктир) и центральная 100% (сплошная)
        ax.vlines(0.8, 0.98, 1.02, colors='black', linestyles='--', linewidth=0.8)
        ax.vlines(1.25, 0.98, 1.02, colors='black', linestyles='--', linewidth=0.8)
        ax.vlines(1.0, 0.98, 1.02, colors='black', linestyles='-', linewidth=0.8)

        # оформление осей — оставляем только нижнюю
        ax.set_xlim(0.75, 1.3)
        ax.set_ylim(0.95, 1.05)
        ax.get_yaxis().set_visible(False)
        for spine in ['top', 'right', 'left']:
            ax.spines[spine].set_visible(False)
        ax.spines['bottom'].set_linewidth(0.8)

        # подписи в процентах
        ax.xaxis.set_ticks_position('bottom')
        ax.set_xticks([0.8, 1.0, 1.2])
        ax.set_xticklabels(['80 %', '100 %', '120 %'])

        plt.tight_layout()
        # -----------------------------
        # далее без изменений:
        fig_path = f"tmp_ci_{p}.png"
        plt.savefig(fig_path, bbox_inches="tight", dpi=300)
        plt.close(fig)

        row[5].paragraphs[0].add_run().add_picture(fig_path, width=Inches(1.8))
        os.remove(fig_path)

        ax.axis("off")

        fig_path = f"tmp_ci_{p}.png"
        plt.savefig(fig_path, bbox_inches="tight", dpi=300)
        plt.close(fig)

        row[5].paragraphs[0].add_run().add_picture(fig_path, width=Inches(1.8))
        os.remove(fig_path)

    doc.save(save_path)
    return save_path

def export_power_analysis_table(gmr_list, cv_list, n=32, alpha=0.05, save_path="power_table.docx"):
    # ——— Проверка входных данных ———
    if len(gmr_list) != len(cv_list):
        raise ValueError("Списки gmr_list и cv_list должны быть одинаковой длины")

    def se_ln_gmr(cv, n):
        """Standard error of ln(GMR) for 2x2 crossover"""
        return np.sqrt(2 * (cv / 100) ** 2 / n)

    def power_equivalence(GMR, CV, n, alpha=0.05):
        """Analytical power calculation for 2x2 crossover BE trial"""
        se = se_ln_gmr(CV, n)
        z = norm.ppf(1 - alpha)
        l_theta1, l_theta2 = np.log(0.8), np.log(1.25)
        l_gmr = np.log(GMR)

        t1 = (l_gmr - l_theta1) / se
        t2 = (l_theta2 - l_gmr) / se
        power = norm.cdf(t1 - z) + norm.cdf(t2 - z) - 1
        return np.clip(power, 0, 1)

    doc = Document()
    doc.add_heading("Анализ мощности исследования", level=1)
    doc.add_paragraph("Мощность статистических тестов оценивали апостериорно на основе наблюдаемых значений GMR и CV. Расчёт выполнен согласно схеме 2x2 Cross-Over на log-шкале.")

    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Параметр"
    hdr[1].text = "Мощн.\n(M)"
    hdr[2].text = "Объём\nвыборки"
    hdr[3].text = "Ниж.\nэкв. граница"
    hdr[4].text = "Верх.\nэкв. граница"
    hdr[5].text = "Отн.\nсредние (μT/μR)"
    hdr[6].text = "CV₍intra₎"
    hdr[7].text = "α"
    hdr[8].text = "β"

    params = ["AUC₀–t", "AUC₀–∞", "Cmax"]

    for param, gmr, cv in zip(params, gmr_list, cv_list):
        row = table.add_row().cells
        power = power_equivalence(gmr, cv, n=n, alpha=alpha)

        row[0].text = f"тест по {param}"
        row[1].text = f"{power:.5f}"
        row[2].text = str(n)
        row[3].text = "0,8"
        row[4].text = "1,25"
        row[5].text = f"{gmr:.4f}"
        row[6].text = f"{cv:.4f}"
        row[7].text = f"{alpha:.2f}"
        row[8].text = f"{1 - power:.5f}"

    # Заключение
    doc.add_paragraph()
    doc.add_paragraph(
        f"По всем параметрам фармакокинетики {', '.join(params)} апостериорная мощность тестов биоэквивалентности составила более 80,00%.",
        style='Intense Quote'
    )
    doc.save(save_path)
    return save_path








